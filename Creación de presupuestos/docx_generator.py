from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import streamlit as st



# PRESUPUESTO/OFERTA ECONÓMICA------------------


def docx_generator(ref,cliente,obra,fecha, *args):

    doc = Document(r"C:\Users\Cash\Downloads\PRESUPUESTO - MODELO - COPIA LOGO.docx")
        
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run =p.add_run("PRESUPUESTO/OFERTA ECONÓMICA")
    print("hola")
    font = run.font
    font.name = 'Aptos Narrow'
    font.size = Pt(12)
    font.bold = True
    font.underline = True

    # ----------------------------------------------
    # Espacio
    doc.add_paragraph()

    # REF-------------------------------------------
    p1 = doc.add_paragraph("")
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    run1  = p1.add_run(f"REF: {ref}")
    font1 = run1.font
    font1.name = 'Verdana'
    font1.size = Pt(12)
    font1.bold = True
    # FECHA-------------------------------------------
    p2 = doc.add_paragraph("")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2  = p2.add_run(f"FECHA: {fecha}")
    font2 = run2.font
    font2.name = 'Verdana'
    font2.size = Pt(12)
    font2.bold = True
    # CLIENTE-------------------------------------------
    p3 = doc.add_paragraph("")
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    run3  = p3.add_run(f"CLIENTE: {cliente}")
    font3 = run3.font
    font3.name = 'Verdana'
    font3.size = Pt(12)
    font3.bold = True
    # OBRA-------------------------------------------

    p4 = doc.add_paragraph("")
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)
    run4  = p4.add_run(f"OBRA: {obra}")
    font4 = run4.font
    font4.name = 'Verdana'
    font4.size = Pt(12)
    font4.bold = True

    p5 = doc.add_paragraph("")
    # Tabla RESUMEN--------------------------------
    table = doc.add_table(rows=1,cols=1)
    cell = table.cell(0, 0)

    # Cambiar color de dentro de la celda-----------
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'D9E1F2')  # Color azul claro en hexadecimal
    cell._tc.get_or_add_tcPr().append(shading_elm)
    p6 = cell.paragraphs[0]
    p6.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run6 = p6.add_run("Resumen")
    font6 = run6.font
    font6.name = 'Aptos Narrow'
    font6.size = Pt(14)
    font6.bold = True #Negrita
    font6.italic = True #Cursiva
    p6.paragraph_format.space_before = Pt(0)
    p6.paragraph_format.space_after = Pt(0)
    # font6.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos Narrow')

    # Aplicación de cambios especiales--------------
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Agregar linea-------------------------------------
    p7 = doc.add_paragraph("")

    # for i in range(10):
    #     globals()[f"x{i}"] = doc.add_paragraph(f"Valor {i}")

    p_counter = 8

    for arg,arg_value in args:
        if arg == "espacio":
            px = doc.add_paragraph("")

        elif arg == "importe":
            px = doc.add_paragraph("")
            px.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run_px = px.add_run(f"TOTAL {arg_value.replace('.',',')}")
            font_px = run_px.font
            font_px.name = 'Aptos Narrow'
            font_px.size = Pt(14)
            font_px.bold = True
            # p_pr_px = px._p.get_or_add_pPr()

            # borders = OxmlElement('w:pBdr')
            # bottom = OxmlElement('w:bottom')
            # bottom.set(qn('w:val'), 'single')
            # bottom.set(qn('w:sz'), '12')      
            # bottom.set(qn('w:space'), '1')          
            # bottom.set(qn('w:color'), '000000')     

            # # Adjuntar el borde inferior
            # borders.append(bottom)
            # p_pr_px.append(borders)
            # px.paragraph_format.space_before = Pt(0)
            # px.paragraph_format.space_after = Pt(0)
            
        elif arg == "borde_separación":
            # px = doc.add_paragraph("")
            # px.paragraph_format.space_before = Pt(0)
            # px.paragraph_format.space_after = Pt(0)
            # px.add_run("valores")
            p_pr_px = px._p.get_or_add_pPr()

            borders = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')      
            bottom.set(qn('w:space'), '1')          
            bottom.set(qn('w:color'), '000000')     

            # Adjuntar el borde inferior
            borders.append(bottom)
            p_pr_px.append(borders)


        elif arg == "linea":
            px = doc.add_paragraph("")
            px.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run_px =px.add_run(f"{arg_value}")
            font_px = run_px.font
            font_px.name = 'Aptos Narrow'
            font_px.size = Pt(14)
            font_px.bold = True

        elif arg == "sub-linea":
            doc.add_paragraph(
            f"• {arg_value}"
            ).paragraph_format.left_indent = Inches(0.20)

        elif arg == "forma_pago":
            px1 = doc.add_paragraph()
            px1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            runx1 =px1.add_run("NOTA:")
            fontx1 = runx1.font
            fontx1.name = 'Aptos Narrow'
            fontx1.size = Pt(14)
            fontx1.bold = False
            px1.paragraph_format.space_before = Pt(0)
            px1.paragraph_format.space_after = Pt(0)
            # ------------------------------------
            px2 = doc.add_paragraph()
            px2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            runx2 =px2.add_run("IVA NO INCLUIDO")
            fontx2 = runx2.font
            fontx2.name = 'Aptos Narrow'
            fontx2.size = Pt(14)
            fontx2.bold = True
            px2.paragraph_format.space_before = Pt(0)
            px2.paragraph_format.space_after = Pt(0)
            # ------------------------------------
            px3 = doc.add_paragraph()
            px3.alignment = WD_ALIGN_PARAGRAPH.LEFT
            runx3 =px3.add_run(f"FORMA DE PAGO: {arg_value}")
            fontx3 = runx3.font
            fontx3.name = 'Aptos Narrow'
            fontx3.size = Pt(14)
            fontx3.bold = True
            px3.paragraph_format.space_before = Pt(0)
            px3.paragraph_format.space_after = Pt(0)





        

        p_counter = p_counter + 1


    # p8 = doc.add_paragraph("")
    # p8.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # run8 =p8.add_run("PISOS + TERRAZZA")
    # font8 = run8.font
    # font8.name = 'Aptos Narrow'
    # font8.size = Pt(14)
    # font8.bold = True


    # # Agregar sublinea(con viñetas)---------------------
    # # Al modificarse un archivo de Word el cual ya esta creado no se puede crear la viñeta
    # doc.add_paragraph(
    #     "• Limpieza de cristales y carpintería metálica."
    # ).paragraph_format.left_indent = Inches(0.20)

    # doc.add_paragraph(
    #     "• Limpieza de los balcones y terrazas de las viviendas."
    # ).paragraph_format.left_indent = Inches(0.20)
    # doc.add_paragraph(
    #     "• Limpieza de mobiliario, grifería y accesorios de cocina."
    # ).paragraph_format.left_indent = Inches(0.20)
    # doc.add_paragraph(
    #     "• Limpieza de sanitarios, espejos, grifería, mobiliario y azulejos de baños y aseo."
    # ).paragraph_format.left_indent = Inches(0.20)
    # doc.add_paragraph(
    #     "• Barrido y fregado de pavimento de viviendas"
    # ).paragraph_format.left_indent = Inches(0.20)
    # # Agregar parrafo y en el parrafo siguiente poner el importe

    # doc.add_paragraph()
    # # Agregar importe total con borde inferior---------
    # p9 = doc.add_paragraph()
    # p9.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # run9 =p9.add_run("TOTAL 6271,65")
    # font9 = run9.font
    # font9.name = 'Aptos Narrow'
    # font9.size = Pt(14)
    # font9.bold = True

    # p_pr = p9._p.get_or_add_pPr()

    # borders = OxmlElement('w:pBdr')
    # bottom = OxmlElement('w:bottom')
    # bottom.set(qn('w:val'), 'single')       # tipo de borde: 'single', 'dashed', 'double', etc.
    # bottom.set(qn('w:sz'), '12')            # tamaño (en 1/8 puntos)
    # bottom.set(qn('w:space'), '1')          # espacio entre texto y borde
    # bottom.set(qn('w:color'), '000000')     # color en hexadecimal

    # # Adjuntar el borde inferior
    # borders.append(bottom)
    # p_pr.append(borders)


    # p10 = doc.add_paragraph()
    # p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # run10 =p10.add_run("NOTA:")
    # font10 = run10.font
    # font10.name = 'Aptos Narrow'
    # font10.size = Pt(14)
    # font10.bold = False
    # p10.paragraph_format.space_before = Pt(0)
    # p10.paragraph_format.space_after = Pt(0)
    # # ------------------------------------
    # p11 = doc.add_paragraph()
    # p11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # run11 =p11.add_run("IVA NO INCLUIDO")
    # font11 = run11.font
    # font11.name = 'Aptos Narrow'
    # font11.size = Pt(14)
    # font11.bold = True
    # p11.paragraph_format.space_before = Pt(0)
    # p11.paragraph_format.space_after = Pt(0)
    # # ------------------------------------
    # p12 = doc.add_paragraph()
    # p12.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # run12 =p12.add_run("FORMA DE PAGO: TRANSFERENCIA")
    # font12 = run12.font
    # font12.name = 'Aptos Narrow'
    # font12.size = Pt(14)
    # font12.bold = True
    # p12.paragraph_format.space_before = Pt(0)
    # p12.paragraph_format.space_after = Pt(0)

    return doc
    # doc.save(r"C:\Users\Cash\Desktop\proyectos julio\streamlit\Creación de presupuestos\budget files\valores - COPIA.docx")
